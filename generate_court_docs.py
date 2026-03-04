#!/usr/bin/env python3
"""
WCD-42 through WCD-45: Generate updated court documents (.docx)
P31 Labs — BONDING evidence system

Generates 4 Word documents:
  - Exhibit-A-System-Description-v3.docx
  - GAL-Briefing-Memo-v3.docx
  - NGSS-Alignment-v3.docx
  - Engagement-Report-Template-v2.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ── Shared styling ──

def set_style(doc):
    """Configure base styles for all court documents."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_bold_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

# ═══════════════════════════════════════════════════════
# WCD-42: Exhibit A System Description v3
# ═══════════════════════════════════════════════════════

def generate_exhibit_a():
    doc = Document()
    set_style(doc)

    # Title
    title = doc.add_heading('Exhibit A: BONDING Application System Description', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    p.add_run('Version 3.0 — March 2026').italic = True
    p.add_run('\nPrepared by William R. Johnson, P31 Labs')

    doc.add_paragraph('─' * 60)

    # Section 1
    add_heading_styled(doc, 'Section 1: What BONDING Is')
    doc.add_paragraph(
        'BONDING is an educational chemistry game designed for children ages six and older. '
        'Players build molecules by dragging atoms onto a three-dimensional canvas and connecting them '
        'according to real chemical bonding rules. The game contains 10 chemical elements and over 60 '
        'known molecules across three difficulty levels designed for different ages.'
    )
    doc.add_paragraph(
        'The game runs in any modern web browser on phones, tablets, and computers. No download, '
        'account, or payment is required. Once loaded, the game works offline.'
    )

    # Section 2
    add_heading_styled(doc, 'Section 2: How It Works')
    doc.add_paragraph(
        'Players select a difficulty mode before starting:'
    )

    items = [
        ('Seed', 'Designed for younger children (ages 6-7). Uses simplified interactions with '
         'large touch targets and basic molecules like water (H\u2082O) and carbon dioxide (CO\u2082).'),
        ('Sprout', 'Designed for children ages 8-10. Introduces more complex molecules, quest '
         'chains, and intermediate-level educational facts.'),
        ('Sapling', 'Designed for older children and advanced learners. Includes scientific-level '
         'facts, complex molecules, and all available elements.'),
    ]
    for name, desc in items:
        p = doc.add_paragraph()
        p.add_run(f'{name}: ').bold = True
        p.add_run(desc)

    doc.add_paragraph(
        'The game includes guided tutorials for first-time players. A step-by-step walkthrough '
        'teaches new players how to drag atoms, form bonds, and complete their first molecule.'
    )
    doc.add_paragraph(
        'Players receive educational facts about each element and molecule as they build. '
        'Facts are presented at the appropriate reading level for each difficulty mode.'
    )
    doc.add_paragraph(
        'Touch controls are designed for children as young as six, with large targets (56 pixels '
        'with expanded touch areas), simplified drag interactions, and protections against accidental '
        'gestures like pull-to-refresh or edge swipes.'
    )
    doc.add_paragraph(
        'Two or more players can play together remotely using a six-character room code. '
        'Each player sees what the other builds in real time and can send reaction pings '
        '(such as a heart or checkmark) to encourage each other.'
    )

    # Section 3
    add_heading_styled(doc, 'Section 3: What It Teaches')
    doc.add_paragraph(
        'BONDING teaches fundamental chemistry concepts through hands-on building:'
    )
    concepts = [
        'Elements are the building blocks of matter',
        'Atoms combine to form molecules following specific rules',
        'Different atoms form different numbers of bonds (valence)',
        'Molecular geometry depends on the atoms involved',
        'Common molecules have names, formulas, and real-world uses',
        'Chemical reactions follow conservation of atoms',
    ]
    for c in concepts:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_paragraph(
        'The educational content is aligned with Next Generation Science Standards (NGSS) for '
        'grades K-8. A separate alignment document maps each game feature to specific NGSS standards.'
    )

    # Section 4
    add_heading_styled(doc, 'Section 4: How Engagement Is Recorded')
    doc.add_paragraph(
        'Every interaction in the game is automatically recorded with a timestamp. '
        'This includes:'
    )
    events = [
        'When a player places an atom',
        'When a bond forms between two atoms',
        'When a molecule is completed',
        'When an achievement is unlocked',
        'When a player sends or receives a ping',
        'When a quest step is completed',
        'When a session starts and ends',
    ]
    for e in events:
        doc.add_paragraph(e, style='List Bullet')

    doc.add_paragraph(
        'This data is stored locally on the device. It can be exported as a plain-text report '
        'showing exactly when each player was active, what they built, and how they interacted '
        'with other players.'
    )
    doc.add_paragraph(
        'A sample log entry looks like this:'
    )
    sample = doc.add_paragraph(
        '2026-03-10 14:23:07 \u2014 molecule_completed \u2014 Player: Will \u2014 '
        'Formula: H\u2082O (Water) \u2014 Atoms: 3 \u2014 Mode: Seed'
    )
    sample.style = doc.styles['Normal']
    for run in sample.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

    doc.add_paragraph(
        'The log cannot be edited within the game. It provides an objective, timestamped record '
        'of every interaction during each play session.'
    )

    # Section 5
    add_heading_styled(doc, 'Section 5: Who Built It and Why')
    doc.add_paragraph(
        'The game was built by a father for his children\u2019s birthdays. William R. Johnson '
        'designed and developed BONDING as a birthday gift for his son Sebastian, who turns 10 on '
        'March 10, 2026. The game is also designed for his daughter Willow, age 6.'
    )
    doc.add_paragraph(
        'BONDING serves as a vehicle for meaningful, documented engagement between a parent and '
        'children during periods of physical separation. Every atom placed, molecule built, and '
        'ping exchanged is a timestamped record of a father playing alongside his kids.'
    )
    doc.add_paragraph(
        'Development began on February 25, 2026, and the game was completed and deployed in '
        'early March 2026. The game has been tested with 470 automated quality checks to ensure '
        'it works reliably on the children\u2019s devices.'
    )

    # Section 6
    add_heading_styled(doc, 'Section 6: Safety and Privacy')
    doc.add_paragraph(
        'BONDING collects no personal information. There are no accounts, no passwords, no email '
        'addresses, and no tracking. The game does not connect to any social media or advertising '
        'networks.'
    )
    doc.add_paragraph(
        'Multiplayer sessions use a six-character room code shared privately between players. '
        'There is no public lobby, no chat with strangers, and no way for unknown players to '
        'join a session.'
    )
    doc.add_paragraph(
        'All engagement data remains on the device unless the user manually exports it.'
    )

    # Section 7
    add_heading_styled(doc, 'Section 7: What the Data Shows')
    doc.add_paragraph(
        'This section will be updated with actual engagement data from sessions beginning '
        'March 10, 2026.'
    )
    doc.add_paragraph(
        'Once play sessions begin, this section will include:'
    )
    data_items = [
        'Total number of sessions played',
        'Duration of each session',
        'Molecules built by each player',
        'Pings exchanged between players',
        'Achievements unlocked',
        'Quest progress',
    ]
    for d in data_items:
        doc.add_paragraph(d, style='List Bullet')

    path = os.path.join(OUTPUT_DIR, 'Exhibit-A-System-Description-v3.docx')
    doc.save(path)
    print(f'  [OK] {path}')


# ═══════════════════════════════════════════════════════
# WCD-43: GAL Briefing Memo v3
# ═══════════════════════════════════════════════════════

def generate_gal_memo():
    doc = Document()
    set_style(doc)

    # Header block
    header_items = [
        ('TO:', 'Guardian Ad Litem'),
        ('FROM:', 'William R. Johnson'),
        ('DATE:', 'March 2026'),
        ('RE:', 'BONDING \u2014 Educational Chemistry Game for Sebastian and Willow'),
    ]
    for label, value in header_items:
        p = doc.add_paragraph()
        run_label = p.add_run(label + ' ')
        run_label.bold = True
        p.add_run(value)

    doc.add_paragraph('─' * 60)

    # Body
    doc.add_paragraph(
        'I am writing to make you aware of BONDING, an educational game I built for my '
        'children. The purpose of this memo is to explain what the game is, how it works, '
        'and why I believe it can serve as a useful tool for documented parental engagement.'
    )

    doc.add_paragraph(
        'BONDING is a chemistry-based building game where players construct molecules by '
        'dragging atoms onto a canvas. It runs in a web browser with no download, account, '
        'or payment required. The game includes three difficulty modes: Seed, designed for '
        'Willow (age 6), offers simplified interactions with large touch targets and basic '
        'molecules. Sprout, designed for Sebastian (turning 10), introduces more complex '
        'molecules, guided quest chains, and intermediate educational facts. A third mode, '
        'Sapling, is available for advanced learners.'
    )

    doc.add_paragraph(
        'The educational content is aligned with 6 Next Generation Science Standards spanning '
        'grades K-8. Willow engages with K-2 standards about observable properties of matter. '
        'Sebastian engages with grade 3-8 standards about atomic composition, molecular models, '
        'and conservation of atoms in chemical reactions.'
    )

    doc.add_paragraph(
        'The game has passed 470 automated quality checks and has been tested on the same '
        'tablet devices the children will use. It includes guided tutorials for first-time '
        'players, age-appropriate educational facts, and touch controls designed for children '
        'as young as six.'
    )

    doc.add_paragraph(
        'Every interaction in the game is automatically logged with a timestamp, including '
        'atoms placed, molecules completed, achievements unlocked, and reaction pings '
        'exchanged between players. This data is stored locally on the device and can be '
        'exported as a plain-text report documenting exactly when each player was active, '
        'what they built, and how they interacted with other players.'
    )

    doc.add_paragraph(
        'The game supports multiplayer, allowing a parent and child to play together remotely '
        'on separate devices using a private room code. No strangers can join. There is no '
        'chat, no personal data collection, and no connection to social media.'
    )

    doc.add_paragraph(
        'I respectfully request the opportunity to use BONDING with my children during '
        'scheduled contact, either in person or remotely on separate devices. I am prepared '
        'to provide engagement reports from each session documenting the duration, content, '
        'and interactions.'
    )

    path = os.path.join(OUTPUT_DIR, 'GAL-Briefing-Memo-v3.docx')
    doc.save(path)
    print(f'  [OK] {path}')


# ═══════════════════════════════════════════════════════
# WCD-44: NGSS Alignment v3
# ═══════════════════════════════════════════════════════

def generate_ngss():
    doc = Document()
    set_style(doc)

    title = doc.add_heading('BONDING \u2014 NGSS Standards Alignment', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    p.add_run('Version 3.0 \u2014 March 2026').italic = True
    p.add_run('\nPrepared by William R. Johnson, P31 Labs')

    doc.add_paragraph('─' * 60)

    # Intro
    doc.add_paragraph(
        'This document maps the educational content in BONDING to specific Next Generation '
        'Science Standards (NGSS). Each standard is linked to the game\u2019s difficulty mode '
        'and the child who will primarily engage with it.'
    )

    # Mode mapping
    add_heading_styled(doc, 'Difficulty Mode \u2192 Grade Band Mapping', level=2)

    mode_table = doc.add_table(rows=4, cols=4)
    mode_table.style = 'Table Grid'
    mode_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ['Mode', 'Grade Band', 'Target Child', 'Focus']
    for i, h in enumerate(headers):
        cell = mode_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    rows_data = [
        ['Seed', 'K-2', 'Willow (age 6)', 'Observable properties, basic building'],
        ['Sprout', '3-5', 'Sebastian (age 10)', 'Particle models, composition, bonding rules'],
        ['Sapling', '6-8', 'Advanced learners', 'Conservation of atoms, reaction modeling'],
    ]
    for i, row_data in enumerate(rows_data):
        for j, val in enumerate(row_data):
            mode_table.rows[i + 1].cells[j].text = val

    doc.add_paragraph()

    # Standards table
    add_heading_styled(doc, 'Standards Alignment', level=2)

    standards = [
        ('2-PS1-1', 'K-2', 'Seed', 'Willow',
         'Plan and conduct an investigation to describe and classify different kinds of '
         'materials by their observable properties.',
         'Players identify elements by color, size, and symbol. Seed mode presents atoms as '
         'distinct objects with observable differences.'),
        ('2-PS1-2', 'K-2', 'Seed', 'Willow',
         'Analyze data obtained from testing different materials to determine which materials '
         'have the properties that are best suited for an intended purpose.',
         'Players learn which atoms combine successfully and which do not, discovering through '
         'trial which elements serve specific bonding purposes.'),
        ('5-PS1-1', '3-5', 'Sprout', 'Sebastian',
         'Develop a model to describe that matter is made of particles too small to be seen.',
         'Players build invisible molecules from visible atom representations. The game models '
         'atomic-scale structures that correspond to real substances like water, methane, and salt.'),
        ('MS-PS1-1', '6-8', 'Sprout / Sapling', 'Sebastian',
         'Develop models to describe the atomic composition of simple molecules and extended '
         'structures.',
         'Players construct molecules following real valence rules. The 3D canvas shows molecular '
         'geometry (tetrahedral, linear, bent) determined by the atoms placed.'),
        ('MS-PS1-2', '6-8', 'Sapling', 'Sebastian',
         'Analyze and interpret data on the properties of substances before and after the '
         'substances interact to determine if a chemical reaction has occurred.',
         'Quest chains guide players through building reactants and products, showing how atoms '
         'rearrange during chemical reactions while being conserved.'),
        ('MS-PS1-5', '6-8', 'Sapling', 'Sebastian',
         'Develop and use a model to describe how the total number of atoms does not change in '
         'a chemical reaction and thus mass is conserved.',
         'The game enforces conservation: every atom placed contributes to the final molecule. '
         'Players see that the same atoms appear in both reactants and products during quests.'),
    ]

    std_table = doc.add_table(rows=len(standards) + 1, cols=6)
    std_table.style = 'Table Grid'
    std_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    std_headers = ['Standard', 'Grade', 'Mode', 'Child', 'Standard Text', 'How BONDING Addresses It']
    for i, h in enumerate(std_headers):
        cell = std_table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True

    for i, s in enumerate(standards):
        for j, val in enumerate(s):
            std_table.rows[i + 1].cells[j].text = val

    doc.add_paragraph()

    # Age mapping
    add_heading_styled(doc, 'Child-Specific Standard Engagement', level=2)

    doc.add_paragraph(
        'Willow (age 6) will engage primarily with standards 2-PS1-1 and 2-PS1-2 through '
        'the Seed difficulty mode. These K-2 standards focus on observable properties of '
        'materials and hands-on investigation \u2014 activities that match her developmental '
        'stage and the simplified game interactions available in Seed mode.'
    )

    doc.add_paragraph(
        'Sebastian (age 9, turning 10) will engage with standards 5-PS1-1, MS-PS1-1, and '
        'MS-PS1-2 through the Sprout difficulty mode. These grade 3-8 standards focus on '
        'modeling atomic composition, understanding molecular structure, and analyzing '
        'chemical reactions \u2014 concepts introduced through the game\u2019s quest chains, '
        'educational facts, and progressively complex molecules.'
    )

    # Crosscutting
    add_heading_styled(doc, 'Crosscutting Concepts', level=2)

    crosscutting = [
        ('Patterns', 'Bonding rules are consistent and predictable. Players learn to recognize '
         'which combinations work and apply those patterns to build new molecules.'),
        ('Cause and Effect', 'Placing an atom in a valid position causes a bond to form. '
         'Placing it in an invalid position is rejected with feedback. Players learn that '
         'specific actions produce specific outcomes.'),
        ('Systems and System Models', 'Each molecule is a system of interconnected atoms. '
         'Players build and observe these systems, learning how components relate to the whole.'),
        ('Structure and Function', 'Molecular geometry (shape) determines molecular identity. '
         'Players see that the same atoms arranged differently produce different molecules.'),
    ]
    for name, desc in crosscutting:
        p = doc.add_paragraph()
        p.add_run(f'{name}: ').bold = True
        p.add_run(desc)

    path = os.path.join(OUTPUT_DIR, 'NGSS-Alignment-v3.docx')
    doc.save(path)
    print(f'  [OK] {path}')


# ═══════════════════════════════════════════════════════
# WCD-45: Engagement Report Template v2
# ═══════════════════════════════════════════════════════

def generate_engagement_template():
    doc = Document()
    set_style(doc)

    title = doc.add_heading('BONDING \u2014 Engagement Report', level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Instructions header (for Will's reference)
    inst_box = doc.add_paragraph()
    inst_box.add_run('INSTRUCTIONS (remove before filing):').bold = True
    doc.add_paragraph(
        'After each play session, export Exhibit A data from the game menu. '
        'Copy the session summary into the appropriate date section below. '
        'Fill in all bracketed fields. Print and file with the court.'
    )
    doc.add_paragraph('─' * 60)

    # Report header
    header_fields = [
        ('Report Period:', '[Start Date] through [End Date]'),
        ('Prepared By:', 'William R. Johnson'),
        ('Prepared For:', 'Guardian Ad Litem / Superior Court of [County], Georgia'),
    ]
    for label, value in header_fields:
        p = doc.add_paragraph()
        p.add_run(label + ' ').bold = True
        p.add_run(value)

    doc.add_paragraph()

    # Summary section
    add_heading_styled(doc, 'Summary', level=2)

    summary_table = doc.add_table(rows=10, cols=2)
    summary_table.style = 'Table Grid'
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    summary_fields = [
        ('Total Sessions', '[    ]'),
        ('Total Play Time', '[    ] minutes'),
        ('Unique Molecules Built', '[    ]'),
        ('Total Pings Exchanged', '[    ]'),
        ('Average Session Duration', '[    ] minutes'),
        ('Difficulty Modes Used', '[Seed / Sprout / Sapling]'),
        ('Achievements Unlocked', '[    ]'),
        ('Longest Session', '[    ] minutes'),
        ('Total Atoms Placed', '[    ]'),
        ('Days With Activity', '[    ]'),
    ]
    for i, (label, value) in enumerate(summary_fields):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value
        for p in summary_table.rows[i].cells[0].paragraphs:
            for r in p.runs:
                r.bold = True

    doc.add_paragraph()

    # Session sections (6 templates)
    for n in range(1, 7):
        add_heading_styled(doc, f'Session {n}', level=2)

        session_fields = [
            ('Date:', '[          ]'),
            ('Participants:', 'William R. Johnson, [Child Name(s)]'),
            ('Mode:', '[Seed / Sprout / Sapling]'),
            ('Duration:', '[    ] minutes'),
            ('Molecules Built:', '[List molecules by name]'),
            ('Pings Exchanged:', '[    ]'),
            ('Achievements Unlocked:', '[List achievements]'),
        ]
        for label, value in session_fields:
            p = doc.add_paragraph()
            p.add_run(label + ' ').bold = True
            p.add_run(value)

        add_bold_para(doc, 'Narrative Summary:')
        p = doc.add_paragraph('[Paste generateCourtSummary() output here]')
        for run in p.runs:
            run.font.color.rgb = RGBColor(128, 128, 128)

        add_bold_para(doc, 'Raw Event Log (first 10 entries):')
        p = doc.add_paragraph('[Paste Exhibit A export here]')
        for run in p.runs:
            run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph('─' * 40)

    # Certification
    add_heading_styled(doc, 'Certification', level=2)
    doc.add_paragraph(
        'I certify that the data in this report was automatically generated by the BONDING '
        'application and has not been manually altered. The timestamps, molecule names, and '
        'interaction counts reflect actual gameplay sessions between the listed participants.'
    )

    doc.add_paragraph()
    sig_fields = [
        'Signature: ________________________________',
        'Printed Name: William R. Johnson',
        'Date: [          ]',
    ]
    for s in sig_fields:
        doc.add_paragraph(s)

    path = os.path.join(OUTPUT_DIR, 'Engagement-Report-Template-v2.docx')
    doc.save(path)
    print(f'  [OK] {path}')


# ═══════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════

if __name__ == '__main__':
    print('Generating court documents (WCD-42 through WCD-45)...\n')
    generate_exhibit_a()
    generate_gal_memo()
    generate_ngss()
    generate_engagement_template()
    print('\nAll 4 documents generated.')
